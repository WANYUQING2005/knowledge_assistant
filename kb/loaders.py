# -*- coding: utf-8 -*-
import os, chardet
from typing import List, Tuple
from langchain_core.documents import Document as LCDocument
from pypdf import PdfReader
from docx import Document as DocxDocument

def read_txt_like(path:str)->str:
    raw = open(path,"rb").read()
    enc = chardet.detect(raw).get("encoding") or "utf-8"
    return raw.decode(enc, errors="ignore")

def load_txt(path)->List[LCDocument]:
    return [LCDocument(page_content=read_txt_like(path),
                       metadata={"source":path,"title":os.path.basename(path),"source_type":"file"})]

def load_md(path)->List[LCDocument]:
    return load_txt(path)

def load_pdf(path)->List[LCDocument]:
    reader = PdfReader(path)
    docs=[]
    for i,page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(LCDocument(page_content=text,
                                   metadata={"source":path,"title":os.path.basename(path),
                                             "page":i+1,"source_type":"pdf"}))
    return docs

def load_docx(path)->List[LCDocument]:
    d = DocxDocument(path)
    text = "\n".join([p.text for p in d.paragraphs])
    return [LCDocument(page_content=text,
                       metadata={"source":path,"title":os.path.basename(path),"source_type":"docx"})]

LOADERS = {
    ".txt": load_txt,
    ".md":  load_md,
    ".pdf": load_pdf,
    ".docx":load_docx,
}

def load_any(path)->List[LCDocument]:
    ext = os.path.splitext(path)[1].lower()
    fn = LOADERS.get(ext)
    if not fn:
        return []
    return fn(path)
